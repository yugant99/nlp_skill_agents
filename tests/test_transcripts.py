from pathlib import Path

from docx import Document

from backend.analysis.transcripts import (
    StudyConfig,
    extract_transcript_text,
    parse_transcript,
)


def test_extract_transcript_text_reads_txt_and_docx(tmp_path: Path) -> None:
    txt_path = tmp_path / "sample.txt"
    txt_path.write_text("vr001_c: Hello\nvr001_p: Hi there", encoding="utf-8")

    docx_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("vr001_c: Hello")
    doc.add_paragraph("vr001_p: Hi there")
    doc.save(docx_path)

    assert extract_transcript_text(txt_path) == "vr001_c: Hello\nvr001_p: Hi there"
    assert extract_transcript_text(docx_path) == "vr001_c: Hello\nvr001_p: Hi there"


def test_parse_transcript_creates_ordered_turns_with_configured_roles() -> None:
    content = """
    vr001_c: Um, look at this picture. [laughter]
    vr001_p: I remember that place.
    vr001_c: Do you remember the year?
    """
    config = StudyConfig(
        participant_id="vr001",
        speaker_prefixes={"caregiver": "vr001_c", "participant": "vr001_p"},
        speaker_labels={"caregiver": "Caregiver", "participant": "Participant"},
    )

    transcript = parse_transcript(content, config, source_filename="vr001.txt")

    assert transcript.source_filename == "vr001.txt"
    assert [turn.role for turn in transcript.turns] == [
        "caregiver",
        "participant",
        "caregiver",
    ]
    assert transcript.turns[0].speaker_label == "Caregiver"
    assert transcript.turns[1].text == "I remember that place."
    assert transcript.turns[2].turn_index == 2


def test_parse_transcript_infers_vr_style_prefixes_when_config_is_partial() -> None:
    content = "vr015_c: Hello there.\nvr015_p: Uh, hi."
    config = StudyConfig(participant_id="vr015")

    transcript = parse_transcript(content, config, source_filename="vr015.txt")

    assert [turn.role for turn in transcript.turns] == ["caregiver", "participant"]
    assert transcript.config.speaker_prefixes == {
        "caregiver": "vr015_c",
        "participant": "vr015_p",
    }

