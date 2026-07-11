import json
import sqlite3
from hashlib import sha256
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from backend.app.main import app
from backend.storage.study_store import StudyWorkspaceStore


def test_health_endpoint() -> None:
    client = TestClient(app)

    response = client.get("/api/health")

    assert response.status_code == 200
    assert response.json() == {"status": "ok", "storage": "local"}


def test_storage_schema_status_reports_applied_migrations(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.get("/api/storage/schema-status")

    assert response.status_code == 200
    payload = response.json()
    assert payload["compatible"] is True
    assert payload["databases"]["analysis_runs"]["current_version"] == 5
    assert [
        migration["name"]
        for migration in payload["databases"]["analysis_runs"]["migrations"]
    ] == [
        "create-base-analysis-runs",
        "add-evidence-identity",
        "add-source-import-identity",
        "add-project-source-lineage",
        "index-analysis-run-history",
    ]
    assert payload["databases"]["evidence_catalog"]["current_version"] == 3
    assert [
        migration["name"]
        for migration in payload["databases"]["evidence_catalog"]["migrations"]
    ] == [
        "create-import-catalog",
        "add-project-source-lineage",
        "index-workspace-history",
    ]


def test_storage_schema_status_rejects_newer_database(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    with sqlite3.connect(tmp_path / "evidence.sqlite3") as connection:
        connection.execute("pragma user_version = 99")
    client = TestClient(app)

    response = client.get("/api/storage/schema-status")

    assert response.status_code == 409
    assert "newer than supported version 3" in response.json()["detail"]


def test_default_skill_pack_endpoint() -> None:
    client = TestClient(app)

    response = client.get("/api/skill-packs/default")

    assert response.status_code == 200
    payload = response.json()
    assert payload["id"] == "default_transcript_metrics"
    assert payload["metrics"] == [
        "base_metrics",
        "lexical_metrics",
        "disfluency_metrics",
    ]


def test_validate_dynamic_skill_pack_endpoint() -> None:
    client = TestClient(app)

    response = client.post(
        "/api/skill-packs/validate",
        json={
            "id": "care_study",
            "name": "Care Study",
            "version": "1.0.0",
            "metrics": ["concept_count_metrics", "cue_inventory_metrics"],
            "speaker_roles": {
                "caregiver": {"label": "Care Partner", "prefixes": ["CG"]},
                "participant": {"label": "Participant", "prefixes": ["P"]},
            },
            "concept_lexicons": {"pain": ["pain", "hurts"]},
            "nonverbal_cues": {"pause": ["pause"]},
        },
    )

    assert response.status_code == 200
    assert response.json() == {
        "valid": True,
        "skill_pack": {
            "id": "care_study",
            "name": "Care Study",
            "version": "1.0.0",
            "metric_ids": ["concept_count_metrics", "cue_inventory_metrics"],
            "speaker_roles": {
                "caregiver": "Care Partner",
                "participant": "Participant",
            },
            "speaker_prefixes": {
                "caregiver": ["CG"],
                "participant": ["P"],
            },
            "disfluency_tokens": [],
            "concept_lexicons": {"pain": ["pain", "hurts"]},
            "nonverbal_cues": {"pause": ["pause"]},
        },
    }


def test_validate_dynamic_skill_pack_endpoint_returns_clear_errors() -> None:
    client = TestClient(app)

    response = client.post(
        "/api/skill-packs/validate",
        json={
            "id": "bad",
            "name": "Bad",
            "version": "1.0.0",
            "metrics": ["not_registered"],
        },
    )

    assert response.status_code == 400
    assert "not_registered" in response.json()["detail"]


def test_validate_skill_pack_text_endpoint_accepts_yaml() -> None:
    client = TestClient(app)

    response = client.post(
        "/api/skill-packs/validate-text",
        json={
            "filename": "study.yaml",
            "content": """
id: yaml_pack
name: YAML Pack
version: 1.0.0
metrics:
  - concept_count_metrics
concept_lexicons:
  pain:
    - pain
    - hurts
nonverbal_cues:
  pause:
    - pause
""".strip(),
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["skill_pack"]["id"] == "yaml_pack"
    assert payload["payload"]["metrics"] == ["concept_count_metrics"]
    assert payload["payload"]["concept_lexicons"] == {"pain": ["pain", "hurts"]}


def test_create_run_from_txt_upload(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.post(
        "/api/runs",
        data={
            "config": json.dumps(
                {
                    "participant_id": "vr009",
                    "selected_metrics": ["base_metrics", "disfluency_metrics"],
                    "disfluency_tokens": ["um"],
                }
            )
        },
        files={
            "file": (
                "vr009.txt",
                b"vr009_c: Um, hello there.\nvr009_p: Hello.",
                "text/plain",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["source_filename"] == "vr009.txt"
    assert payload["import_id"].startswith("imp_")
    assert payload["project_source_id"].startswith("psrc_")
    assert payload["parent_transcript_revision_id"] == ""
    assert payload["workspace_id"] == "local-default"
    assert payload["source_blob_sha256"] == sha256(
        b"vr009_c: Um, hello there.\nvr009_p: Hello."
    ).hexdigest()
    assert payload["source_media_type"] == "text/plain"
    assert payload["source_id"].startswith("src_")
    assert len(payload["transcript_sha256"]) == 64
    assert payload["transcript_revision_id"].startswith("trv_")
    assert [result["metric_id"] for result in payload["results"]] == [
        "base_metrics",
        "disfluency_metrics",
    ]
    assert payload["diagnostics"] == {
        "turn_counts": {"caregiver": 1, "participant": 1},
        "warnings": [],
    }
    assert payload["stored"]["results_json"].endswith("results.json")
    assert payload["exports"] == [
        {
            "metric_id": "base_metrics",
            "filename": "base_metrics.csv",
            "download_url": f"/api/runs/{payload['run_id']}/exports/base_metrics.csv",
        },
        {
            "metric_id": "disfluency_metrics",
            "filename": "disfluency_metrics.csv",
            "download_url": f"/api/runs/{payload['run_id']}/exports/disfluency_metrics.csv",
        },
    ]
    assert (tmp_path / "runs.sqlite3").exists()
    assert (tmp_path / "evidence.sqlite3").exists()

    imports_response = client.get("/api/evidence/imports")
    assert imports_response.status_code == 200
    assert imports_response.json()["imports"][0]["import_id"] == payload["import_id"]
    blob_response = client.get(
        f"/api/evidence/blobs/{payload['source_blob_sha256']}/verify"
    )
    assert blob_response.status_code == 200
    assert blob_response.json() == {
        "source_blob_sha256": payload["source_blob_sha256"],
        "verified": True,
        "size_bytes": len(b"vr009_c: Um, hello there.\nvr009_p: Hello."),
    }


def test_download_export_csv_for_run(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    create_response = client.post(
        "/api/runs",
        data={
            "config": json.dumps(
                {
                    "participant_id": "vr010",
                    "selected_metrics": ["base_metrics"],
                }
            )
        },
        files={
            "file": (
                "vr010.txt",
                b"vr010_c: Hello there.\nvr010_p: Hello.",
                "text/plain",
            )
        },
    )
    run_id = create_response.json()["run_id"]

    response = client.get(f"/api/runs/{run_id}/exports/base_metrics.csv")

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/csv")
    assert "attachment; filename=\"base_metrics.csv\"" in response.headers[
        "content-disposition"
    ]
    assert response.text.startswith("speaker,turns,clean_words")
    assert "caregiver,1,2" in response.text


def test_download_export_rejects_path_traversal(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.get("/api/runs/abc/exports/../runs.sqlite3")

    assert response.status_code == 404


def test_list_runs_returns_recent_local_runs(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    for filename, participant in [("one.txt", "vr040"), ("two.txt", "vr041")]:
        client.post(
            "/api/runs",
            data={
                "config": json.dumps(
                    {
                        "participant_id": participant,
                        "selected_metrics": ["base_metrics"],
                    }
                )
            },
            files={
                "file": (
                    filename,
                    f"{participant}_c: Hello.\n{participant}_p: Hi.".encode(),
                    "text/plain",
                )
            },
        )

    response = client.get("/api/runs")

    assert response.status_code == 200
    assert [row["source_filename"] for row in response.json()["runs"]] == [
        "two.txt",
        "one.txt",
    ]


def test_create_run_surfaces_diagnostic_warnings(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.post(
        "/api/runs",
        data={
            "config": json.dumps(
                {
                    "participant_id": "vr099",
                    "selected_metrics": ["base_metrics"],
                }
            )
        },
        files={
            "file": (
                "bad.txt",
                b"This transcript has no known speaker prefixes.",
                "text/plain",
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["diagnostics"]["warnings"] == [
        {
            "code": "no_turns_found",
            "message": "No speaker turns were detected. Check participant ID and speaker prefixes.",
        }
    ]


def test_create_run_accepts_custom_speaker_prefixes(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.post(
        "/api/runs",
        data={
            "config": json.dumps(
                {
                    "participant_id": "dyad01",
                    "speaker_prefixes": {
                        "caregiver": "care_partner",
                        "participant": "participant",
                    },
                    "selected_metrics": ["base_metrics"],
                }
            )
        },
        files={
            "file": (
                "dyad01.txt",
                b"care_partner: Hello there.\nparticipant: Hello back.",
                "text/plain",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["turn_count"] == 2
    assert payload["diagnostics"]["turn_counts"] == {
        "caregiver": 1,
        "participant": 1,
    }
    assert payload["results"][0]["rows"][0]["clean_words"] == 2


def test_create_run_from_text_payload(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.post(
        "/api/runs/text",
        json={
            "source_filename": "pasted_transcript.txt",
            "content": "vr050_c: Hello.\nvr050_p: Um, hello back.",
            "config": {
                "participant_id": "vr050",
                "selected_metrics": ["base_metrics", "disfluency_metrics"],
            },
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["source_filename"] == "pasted_transcript.txt"
    assert payload["turn_count"] == 2
    assert payload["results"][1]["rows"][-1]["disfluency_count"] == 1


def test_text_run_api_records_and_validates_revision_lineage(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    config = {"participant_id": "vr051", "selected_metrics": ["base_metrics"]}

    first_response = client.post(
        "/api/runs/text",
        json={
            "source_filename": "session.txt",
            "content": "vr051_c: First prompt.\nvr051_p: First response.",
            "config": config,
        },
    )
    first = first_response.json()
    second_response = client.post(
        "/api/runs/text",
        json={
            "source_filename": "session-revised.txt",
            "content": "vr051_c: Revised prompt.\nvr051_p: Revised response.",
            "config": config,
            "project_source_id": first["project_source_id"],
            "parent_transcript_revision_id": first["transcript_revision_id"],
        },
    )
    second = second_response.json()

    assert first_response.status_code == 200
    assert second_response.status_code == 200
    assert second["project_source_id"] == first["project_source_id"]
    assert (
        second["parent_transcript_revision_id"]
        == first["transcript_revision_id"]
    )
    assert second["transcript_revision_id"] != first["transcript_revision_id"]

    history_response = client.get(
        f"/api/evidence/sources/{first['project_source_id']}"
    )
    history = history_response.json()
    assert history_response.status_code == 200
    assert [item["transcript_revision_id"] for item in history["revisions"]] == [
        first["transcript_revision_id"],
        second["transcript_revision_id"],
    ]

    invalid_response = client.post(
        "/api/runs/text",
        json={
            "source_filename": "invalid-revision.txt",
            "content": "vr051_c: Invalid.\nvr051_p: Invalid.",
            "config": config,
            "project_source_id": first["project_source_id"],
            "parent_transcript_revision_id": "trv_missing",
        },
    )
    assert invalid_response.status_code == 400
    assert "Parent revision does not belong" in invalid_response.json()["detail"]

    rootless_response = client.post(
        "/api/runs/text",
        json={
            "source_filename": "rootless-revision.txt",
            "content": "vr051_c: Rootless.\nvr051_p: Rootless.",
            "config": config,
            "project_source_id": first["project_source_id"],
        },
    )
    assert rootless_response.status_code == 400
    assert "existing source requires a parent" in rootless_response.json()["detail"]
    assert len(client.get("/api/evidence/imports").json()["imports"]) == 2
    assert len(list((tmp_path / "runs").glob("*/results.json"))) == 2


def test_create_text_run_applies_embedded_dynamic_skill_pack(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    skill_pack = {
        "id": "care_study",
        "name": "Care Study",
        "version": "1.0.0",
        "metrics": ["concept_count_metrics", "cue_inventory_metrics"],
        "speaker_roles": {
            "caregiver": {"label": "Care Partner", "prefixes": ["CG"]},
            "participant": {"label": "Participant", "prefixes": ["P"]},
        },
        "disfluency_tokens": ["um"],
        "concept_lexicons": {"pain": ["pain", "hurt", "hurts"]},
        "nonverbal_cues": {"pause": ["pause"]},
    }

    response = client.post(
        "/api/runs/text",
        json={
            "source_filename": "care_study.txt",
            "content": "CG: Does it hurt? [pause]\nP: Um, the pain hurts.",
            "config": {"skill_pack": skill_pack},
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["skill_pack"] == {
        "id": "care_study",
        "name": "Care Study",
        "version": "1.0.0",
    }
    assert payload["diagnostics"]["turn_counts"] == {
        "caregiver": 1,
        "participant": 1,
    }
    assert [result["metric_id"] for result in payload["results"]] == [
        "concept_count_metrics",
        "cue_inventory_metrics",
    ]
    assert payload["results"][0]["rows"][0]["match_count"] == 3
    results_json = tmp_path / "runs" / payload["run_id"] / "results.json"
    assert json.loads(results_json.read_text())["skill_pack"] == payload["skill_pack"]


def test_create_run_rejects_unknown_metric_with_400(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.post(
        "/api/runs/text",
        json={
            "source_filename": "bad_metric.txt",
            "content": "vr060_c: Hello.\nvr060_p: Hi.",
            "config": {
                "participant_id": "vr060",
                "selected_metrics": ["not_a_metric"],
            },
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "Unknown metric skill: not_a_metric"


def test_study_workspace_batch_api_creates_aggregate_outputs(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    study_response = client.post(
        "/api/studies",
        json={
            "name": "Question Study",
            "description": "Prompting style across transcripts.",
        },
    )

    assert study_response.status_code == 200
    study_id = study_response.json()["study"]["id"]

    pack_response = client.post(
        f"/api/studies/{study_id}/skill-pack-versions",
        json={
            "id": "question_pack",
            "name": "Question Pack",
            "version": "1.0.0",
            "metrics": ["question_type_metrics"],
            "speaker_roles": {
                "caregiver": {"label": "Caregiver", "prefixes": ["CG"]},
                "participant": {"label": "Participant", "prefixes": ["P"]},
            },
        },
    )

    assert pack_response.status_code == 200
    version_id = pack_response.json()["version"]["version_id"]

    batch_response = client.post(
        f"/api/studies/{study_id}/batches/text",
        json={
            "skill_pack_version_id": version_id,
            "transcripts": [
                {
                    "source_filename": "one.txt",
                    "metadata": {
                        "participant_id": "P1",
                        "condition": "home",
                        "week": "week_1",
                    },
                    "content": "CG: How are you?\nP: Fine.",
                },
                {
                    "source_filename": "two.txt",
                    "metadata": {
                        "participant_id": "P2",
                        "condition": "lab",
                        "week": "week_1",
                    },
                    "content": "CG: Did sleep improve?\nP: Yes.",
                },
            ],
        },
    )

    assert batch_response.status_code == 200
    payload = batch_response.json()
    assert payload["batch"]["run_count"] == 2
    assert payload["batch"]["failure_count"] == 0
    assert payload["aggregate_results_json"].endswith("aggregate_results.json")
    assert payload["results"][0]["metric_id"] == "question_type_metrics"
    assert payload["results"][0]["rows"][0]["participant_id"] == "P1"
    assert payload["results"][0]["rows"][0]["condition"] == "home"
    assert payload["results"][0]["rows"][0]["week"] == "week_1"
    assert payload["results"][0]["rows"][3]["participant_id"] == "P2"
    assert payload["results"][0]["rows"][3]["condition"] == "lab"
    assert payload["exports"] == [
        {
            "metric_id": "question_type_metrics",
            "filename": "question_type_metrics.csv",
            "path": f"{payload['batch']['aggregate_dir']}/question_type_metrics.csv",
        }
    ]

    list_response = client.get("/api/studies")
    assert list_response.status_code == 200
    assert list_response.json()["studies"][0]["id"] == "question-study"

    bundle_response = client.post(f"/api/studies/{study_id}/bundle")

    assert bundle_response.status_code == 200
    bundle_payload = bundle_response.json()["bundle"]
    assert bundle_payload["study_id"] == "question-study"
    assert bundle_payload["manifest_path"].endswith("manifest.json")

    audit_response = client.get("/api/audit-events")

    assert audit_response.status_code == 200
    event_types = [event["event_type"] for event in audit_response.json()["events"]]
    assert event_types[-4:] == [
        "study.created",
        "skill_pack.versioned",
        "batch.completed",
        "bundle.exported",
    ]


def test_study_workspace_file_batch_api_accepts_txt_and_docx(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    study_response = client.post(
        "/api/studies",
        json={"name": "Multipart Study"},
    )
    study_id = study_response.json()["study"]["id"]
    pack_response = client.post(
        f"/api/studies/{study_id}/skill-pack-versions",
        json={
            "id": "multipart_pack",
            "name": "Multipart Pack",
            "version": "1.0.0",
            "metrics": ["base_metrics"],
        },
    )
    version_id = pack_response.json()["version"]["version_id"]

    docx_buffer = BytesIO()
    doc = Document()
    doc.add_paragraph("P2_c: Did balance improve?")
    doc.add_paragraph("P2_p: It improved a little.")
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    docx_bytes = docx_buffer.getvalue()

    response = client.post(
        f"/api/studies/{study_id}/batches/files",
        data={
            "skill_pack_version_id": version_id,
            "metadata": json.dumps(
                {
                    "P1_home_week1.txt": {
                        "participant_id": "P1",
                        "condition": "home",
                        "week": "week_1",
                    },
                    "P2_lab_week2.docx": {
                        "participant_id": "P2",
                        "condition": "lab",
                        "week": "week_2",
                    },
                }
            ),
        },
        files=[
            (
                "files",
                (
                    "P1_home_week1.txt",
                    b"P1_c: How did walking feel?\nP1_p: It felt steady.",
                    "text/plain",
                ),
            ),
            (
                "files",
                (
                    "P2_lab_week2.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
        ],
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["batch"]["run_count"] == 2
    assert payload["batch"]["failure_count"] == 0
    rows = payload["results"][0]["rows"]
    assert rows[0]["participant_id"] == "P1"
    assert rows[0]["condition"] == "home"
    assert rows[0]["week"] == "week_1"
    assert rows[0]["turns"] == 1
    assert rows[3]["participant_id"] == "P2"
    assert rows[3]["condition"] == "lab"
    assert rows[3]["week"] == "week_2"
    assert rows[3]["turns"] == 1
    imports = client.get("/api/evidence/imports").json()["imports"]
    docx_import = next(
        item for item in imports if item["source_filename"] == "P2_lab_week2.docx"
    )
    assert docx_import["source_blob_sha256"] == sha256(docx_bytes).hexdigest()
    verify_response = client.get(
        f"/api/evidence/blobs/{docx_import['source_blob_sha256']}/verify"
    )
    assert verify_response.json()["size_bytes"] == len(docx_bytes)


def test_study_backup_and_restore_api_round_trips_project(tmp_path, monkeypatch) -> None:
    source_root = tmp_path / "source"
    restore_root = tmp_path / "restore"
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(source_root))
    client = TestClient(app)
    study_response = client.post("/api/studies", json={"name": "Backup API Study"})
    study_id = study_response.json()["study"]["id"]
    version_response = client.post(
        f"/api/studies/{study_id}/skill-pack-versions",
        json={
            "id": "backup_api_pack",
            "name": "Backup API Pack",
            "version": "1.0.0",
            "metrics": ["base_metrics"],
        },
    )
    client.post(
        f"/api/studies/{study_id}/batches/text",
        json={
            "skill_pack_version_id": version_response.json()["version"][
                "version_id"
            ],
            "transcripts": [
                {
                    "source_filename": "session.txt",
                    "content": "P1_c: One.\nP1_p: Two.",
                }
            ],
        },
    )

    backup_response = client.post(f"/api/studies/{study_id}/backup")
    backup = backup_response.json()["backup"]
    archive_bytes = Path(backup["archive_path"]).read_bytes()

    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(restore_root))
    restore_response = client.post(
        "/api/studies/restore",
        files={"file": ("backup.nlpstudy.zip", archive_bytes, "application/zip")},
    )

    assert backup_response.status_code == 200
    assert len(backup["archive_sha256"]) == 64
    assert restore_response.status_code == 200
    assert restore_response.json()["restore"]["study_id"] == study_id
    assert restore_response.json()["restore"]["audit_event_count"] == 3
    assert client.get("/api/studies").json()["studies"][0]["id"] == study_id
    restored_import = client.get("/api/evidence/imports").json()["imports"][0]
    assert client.get(
        f"/api/evidence/blobs/{restored_import['source_blob_sha256']}/verify"
    ).status_code == 200
    conflict_response = client.post(
        "/api/studies/restore",
        files={"file": ("backup.nlpstudy.zip", archive_bytes, "application/zip")},
    )
    assert conflict_response.status_code == 409


def test_study_schema_api_persists_casebook_design(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    study_response = client.post(
        "/api/studies",
        json={"name": "Schema API Study"},
    )
    study_id = study_response.json()["study"]["id"]

    update_response = client.put(
        f"/api/studies/{study_id}/schema",
        json={
            "participant_count": 8,
            "conditions": ["home", "lab", "clinic"],
            "week_count": 3,
            "custom_fields": ["site", "arm"],
        },
    )
    read_response = client.get(f"/api/studies/{study_id}/schema")

    assert update_response.status_code == 200
    schema = update_response.json()["schema"]
    assert schema["study_id"] == study_id
    assert schema["participants"] == [
        "P1",
        "P2",
        "P3",
        "P4",
        "P5",
        "P6",
        "P7",
        "P8",
    ]
    assert schema["conditions"] == ["home", "lab", "clinic"]
    assert schema["weeks"] == ["week_1", "week_2", "week_3"]
    assert schema["custom_fields"] == ["site", "arm"]
    assert read_response.status_code == 200
    assert read_response.json()["schema"] == schema

    oversized_response = client.put(
        f"/api/studies/{study_id}/schema",
        json={"participant_count": 10_001},
    )
    assert oversized_response.status_code == 422


def test_study_batch_history_api_lists_and_loads_results(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    study_response = client.post("/api/studies", json={"name": "History API Study"})
    study_id = study_response.json()["study"]["id"]
    pack_response = client.post(
        f"/api/studies/{study_id}/skill-pack-versions",
        json={
            "id": "history_api_pack",
            "name": "History API Pack",
            "version": "1.0.0",
            "metrics": ["base_metrics"],
        },
    )
    version_id = pack_response.json()["version"]["version_id"]
    first_response = client.post(
        f"/api/studies/{study_id}/batches/text",
        json={
            "skill_pack_version_id": version_id,
            "transcripts": [{"source_filename": "one.txt", "content": "CG: Hello.\nP: Hi."}],
        },
    )
    second_response = client.post(
        f"/api/studies/{study_id}/batches/text",
        json={
            "skill_pack_version_id": version_id,
            "transcripts": [{"source_filename": "two.txt", "content": "CG: Again?\nP: Yes."}],
        },
    )

    list_response = client.get(f"/api/studies/{study_id}/batches")
    loaded_response = client.get(
        f"/api/studies/{study_id}/batches/{first_response.json()['batch']['batch_id']}"
    )

    assert list_response.status_code == 200
    batch_ids = [batch["batch_id"] for batch in list_response.json()["batches"]]
    assert batch_ids == [
        second_response.json()["batch"]["batch_id"],
        first_response.json()["batch"]["batch_id"],
    ]
    assert loaded_response.status_code == 200
    loaded = loaded_response.json()
    assert loaded["batch"]["batch_id"] == first_response.json()["batch"]["batch_id"]
    assert loaded["results"][0]["metric_id"] == "base_metrics"
    assert loaded["results"][0]["rows"][0]["source_filename"] == "one.txt"


def test_study_batch_run_drilldown_api_lists_and_loads_one_run(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    study_response = client.post("/api/studies", json={"name": "Run Drilldown API Study"})
    study_id = study_response.json()["study"]["id"]
    pack_response = client.post(
        f"/api/studies/{study_id}/skill-pack-versions",
        json={
            "id": "run_drilldown_api_pack",
            "name": "Run Drilldown API Pack",
            "version": "1.0.0",
            "metrics": ["base_metrics"],
        },
    )
    batch_response = client.post(
        f"/api/studies/{study_id}/batches/text",
        json={
            "skill_pack_version_id": pack_response.json()["version"]["version_id"],
            "transcripts": [
                {
                    "source_filename": "P1_home_week1.txt",
                    "metadata": {"participant_id": "P1", "condition": "home", "week": "week_1"},
                    "content": "P1_c: Hello?\nP1_p: Hi.",
                },
                {
                    "source_filename": "P2_lab_week1.txt",
                    "metadata": {"participant_id": "P2", "condition": "lab", "week": "week_1"},
                    "content": "P2_c: Again?\nP2_p: Yes.",
                },
            ],
        },
    )
    batch_id = batch_response.json()["batch"]["batch_id"]

    list_response = client.get(f"/api/studies/{study_id}/batches/{batch_id}/runs")
    run_id = list_response.json()["runs"][0]["run_id"]
    loaded_response = client.get(
        f"/api/studies/{study_id}/batches/{batch_id}/runs/{run_id}"
    )

    assert list_response.status_code == 200
    assert [run["source_filename"] for run in list_response.json()["runs"]] == [
        "P1_home_week1.txt",
        "P2_lab_week1.txt",
    ]
    assert list_response.json()["runs"][0]["metadata"]["participant_id"] == "P1"
    assert loaded_response.status_code == 200
    loaded = loaded_response.json()["run"]
    assert loaded["source_filename"] == "P1_home_week1.txt"
    assert list_response.json()["runs"][0]["source_id"] == loaded["source_id"]
    assert (
        list_response.json()["runs"][0]["transcript_revision_id"]
        == loaded["transcript_revision_id"]
    )
    assert {
        key: value
        for key, value in loaded["turns"][0].items()
        if key != "passage_id"
    } == {
        "turn_index": 0,
        "role": "caregiver",
        "speaker_label": "Caregiver",
        "raw_prefix": "P1_c",
        "text": "Hello?",
    }
    assert loaded["turns"][0]["passage_id"].startswith("psg_")
    assert loaded["results"][0]["metric_id"] == "base_metrics"


def test_study_batch_api_includes_failed_file_details(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    store = StudyWorkspaceStore(tmp_path)
    study = store.create_study({"name": "Failure API Study"})
    version = store.add_skill_pack_version(
        study.id,
        {
            "id": "bad_metric_pack",
            "name": "Bad Metric Pack",
            "version": "1.0.0",
            "metrics": ["not_registered"],
        },
        validate=False,
    )
    batch = store.run_text_batch(
        study.id,
        version.version_id,
        [{"source_filename": "bad.txt", "content": "CG: Hello."}],
    )
    client = TestClient(app)

    response = client.get(f"/api/studies/{study.id}/batches/{batch.batch_id}")

    assert response.status_code == 200
    payload = response.json()
    assert payload["batch"]["failure_count"] == 1
    assert payload["failures"] == [
        {
            "source_filename": "bad.txt",
            "error": "Skill pack references unknown metric id(s): not_registered",
        }
    ]


def test_library_approval_api_records_entries_and_audit(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.post(
        "/api/library/skill-packs",
        json={
            "payload": {
                "id": "approved_pack",
                "name": "Approved Pack",
                "version": "1.0.0",
                "metrics": ["base_metrics"],
            },
            "reviewer": "professor",
            "notes": "Ready for reuse.",
        },
    )

    assert response.status_code == 200
    assert response.json()["entry"]["id"] == "approved_pack"

    list_response = client.get("/api/library")

    assert list_response.status_code == 200
    assert list_response.json()["entries"][0]["entry_type"] == "skill_pack"

    audit_response = client.get("/api/audit-events")
    assert audit_response.json()["events"][-1]["event_type"] == (
        "library.skill_pack.approved"
    )


def test_deployment_profile_endpoint_reports_secure_offline_status(
    tmp_path,
    monkeypatch,
) -> None:
    from backend.llm import openrouter

    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    monkeypatch.delenv("OPENROUTER_API_KEY", raising=False)
    monkeypatch.setattr(openrouter, "_DOTENV_LOADED", True)
    client = TestClient(app)

    response = client.get("/api/deployment-profile/secure-offline")

    assert response.status_code == 200
    assert response.json()["ready"] is True
    assert response.json()["checks"][1]["id"] == "network_llm_disabled"


def test_segmentation_api_lists_and_returns_synthetic_cases() -> None:
    client = TestClient(app)

    list_response = client.get("/api/segmentation/cases")
    case_response = client.get("/api/segmentation/cases/pause_overlap_repair")

    assert list_response.status_code == 200
    cases = list_response.json()["cases"]
    assert [case["case_id"] for case in cases] == [
        "pause_overlap_repair",
        "redaction_omission_nonverbal",
    ]
    assert cases[0]["source"] == "synthetic"
    assert cases[0]["forbidden_source_tokens"] == []

    assert case_response.status_code == 200
    payload = case_response.json()["case"]
    assert payload["case_id"] == "pause_overlap_repair"
    assert "[00:00:00]" in payload["descript_text"]
    assert "([FP])" in payload["gold_text"]


def test_segmentation_api_evaluates_draft_against_synthetic_rules() -> None:
    client = TestClient(app)
    case = client.get("/api/segmentation/cases/redaction_omission_nonverbal").json()[
        "case"
    ]

    good_response = client.post(
        "/api/segmentation/evaluate",
        json={
            "case_id": "redaction_omission_nonverbal",
            "draft_text": case["gold_text"],
        },
    )
    bad_response = client.post(
        "/api/segmentation/evaluate",
        json={
            "case_id": "redaction_omission_nonverbal",
            "draft_text": "P: I saw Nala [redacted]",
        },
    )

    assert good_response.status_code == 200
    good_evaluation = good_response.json()["evaluation"]
    assert good_evaluation["passed_rule_count"] == good_evaluation[
        "configured_rule_count"
    ]
    assert "score" not in good_evaluation
    assert good_evaluation["failures"] == []

    assert bad_response.status_code == 200
    failures = {
        failure["rule_id"]
        for failure in bad_response.json()["evaluation"]["failures"]
    }
    assert failures >= {"redaction-comments", "official-source-guard"}


def test_segmentation_api_rejects_unknown_synthetic_case() -> None:
    client = TestClient(app)

    response = client.get("/api/segmentation/cases/not-real")
    evaluate_response = client.post(
        "/api/segmentation/evaluate",
        json={"case_id": "not-real", "draft_text": "P: Hello."},
    )

    assert response.status_code == 404
    assert evaluate_response.status_code == 404


def test_segmentation_run_api_creates_fetches_and_verifies_rule_specialist_run(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "session.txt",
            "descript_text": "[00:00:00] P: Good morning.\n[00:00:03] Av: Uh yes.",
            "rule_ids": [
                "speaker-markers",
                "timestamp-markers",
                "pause-markers",
                "filled-pauses",
            ],
        },
    )

    assert response.status_code == 200
    run = response.json()["run"]
    assert run["source"] == "researcher_provided"
    assert run["import_id"].startswith("imp_")
    assert run["project_source_id"].startswith("psrc_")
    assert run["parent_transcript_revision_id"] == ""
    assert run["workspace_id"] == "local-default"
    assert len(run["source_blob_sha256"]) == 64
    assert run["source_media_type"] == "text/plain"
    assert run["source_id"].startswith("src_")
    assert len(run["transcript_sha256"]) == 64
    assert run["transcript_revision_id"].startswith("trv_")
    assert run["events"][0]["passage_id"].startswith("psg_")
    assert run["cunit_adjudication"]["decisions"][0]["cunit_ids"]
    assert run["merged_draft"].startswith(
        "Researcher-provided transcript: session"
    )
    assert run["status"] == "verified"
    assert run["rule_plan"][0]["specialist_id"] == "speaker_turn"
    assert run["specialist_outputs"][0]["patches"]
    assert run["specialist_outputs"][0]["evidence"]["artifact_path"].endswith(
        "specialists/speaker_turn.html"
    )

    fetch_response = client.get(f"/api/segmentation/runs/{run['run_id']}")

    assert fetch_response.status_code == 200
    assert fetch_response.json()["run"]["run_id"] == run["run_id"]
    assert fetch_response.json()["run"]["source"] == "researcher_provided"

    verify_response = client.post(f"/api/segmentation/runs/{run['run_id']}/verify")

    assert verify_response.status_code == 200
    assert verify_response.json()["run"]["run_id"] == run["run_id"]
    assert verify_response.json()["run"]["source"] == "researcher_provided"
    evaluation = verify_response.json()["run"]["evaluation"]
    assert evaluation["passed_rule_count"] == evaluation["configured_rule_count"]


def test_segmentation_run_api_accepts_uploaded_txt_file(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    source_bytes = b"[00:00:00] P: Good morning.\n[00:00:03] Av: Uh yes.\n"
    response = client.post(
        "/api/segmentation/runs/files",
        data={
            "rule_ids": json.dumps(
                [
                    "speaker-markers",
                    "timestamp-markers",
                    "pause-markers",
                    "filled-pauses",
                ]
            )
        },
        files={
            "file": (
                "descript_export.txt",
                source_bytes,
                "text/plain",
            )
        },
    )

    assert response.status_code == 200
    run = response.json()["run"]
    assert run["source_filename"] == "descript_export.txt"
    assert run["source"] == "researcher_provided"
    assert run["source_blob_sha256"] == sha256(source_bytes).hexdigest()
    assert run["source_media_type"] == "text/plain"
    assert run["merged_draft"].startswith(
        "Researcher-provided transcript: descript_export"
    )
    assert run["status"] == "verified"
    assert run["events"][0]["source_filename"] == "descript_export.txt"
    blob_response = client.get(
        f"/api/evidence/blobs/{run['source_blob_sha256']}/verify"
    )
    assert blob_response.status_code == 200
    assert blob_response.json()["size_bytes"] == len(source_bytes)


def test_segmentation_run_file_api_rejects_non_txt_upload(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    response = client.post(
        "/api/segmentation/runs/files",
        data={"rule_ids": json.dumps(["speaker-markers"])},
        files={
            "file": (
                "descript_export.docx",
                b"not really a docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "Only TXT segmentation uploads are supported"


def test_segmentation_run_api_rejects_invalid_input(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    empty_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "empty.txt",
            "descript_text": "   ",
            "rule_ids": ["speaker-markers"],
        },
    )
    unknown_rule_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "session.txt",
            "descript_text": "[00:00:00] P: Good morning.",
            "rule_ids": ["not-a-rule"],
        },
    )
    unknown_source_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "session.txt",
            "descript_text": "[00:00:00] P: Good morning.",
            "rule_ids": ["speaker-markers"],
            "source": "external",
        },
    )

    assert empty_response.status_code == 400
    assert "descript_text" in empty_response.json()["detail"]
    assert unknown_rule_response.status_code == 400
    assert "not-a-rule" in unknown_rule_response.json()["detail"]
    assert unknown_source_response.status_code == 422


def test_segmentation_rulebook_api_exposes_coverage_and_limits() -> None:
    client = TestClient(app)

    response = client.get("/api/segmentation/rulebook")

    assert response.status_code == 200
    payload = response.json()["rulebook"]
    assert payload["implemented_rule_count"] == 10
    assert payload["tracked_fixture_rule_count"] == 9
    assert payload["generated_fixture_rule_count"] == 10
    assert payload["validation"]["status"] == "not_domain_validated"
    assert "not accuracy" in payload["validation"]["claim_boundary"]
    assert payload["rule_definitions"][0]["rule_id"] == "speaker-markers"
    assert any(
        area["area_id"] == "cunit-boundaries"
        and area["status"] == "implemented-unvalidated"
        for area in payload["method_areas"]
    )


def test_segmentation_run_rewrite_job_uses_failed_rule_routing(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    create_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "needs_rewrite.txt",
            "descript_text": "[00:00:00] P: Good morning.",
            "rule_ids": ["speaker-markers", "overlap-markers"],
        },
    )
    run = create_response.json()["run"]

    assert run["status"] == "needs_rewrite"
    assert run["cunit_adjudication"]["counted_cunit_count"] == 1
    assert run["cunit_adjudication"]["decisions"][0]["boundary_type"]
    assert run["failure_routes"] == [
        {
            "rule_id": "overlap-markers",
            "specialist_id": "repair_overlap",
            "message": "Expected overlapping speech to be marked with angle brackets.",
        }
    ]

    rewrite_response = client.post(
        f"/api/segmentation/runs/{run['run_id']}/rewrite-job"
    )

    assert rewrite_response.status_code == 200
    payload = rewrite_response.json()
    assert payload["job"]["id"] == f"rewrite_{run['run_id']}"
    assert payload["job"]["source_request_id"] == run["run_id"]
    prompt_path = tmp_path / "agent_jobs" / payload["job"]["id"] / "rewrite_prompt.html"
    prompt = prompt_path.read_text(encoding="utf-8")
    assert "overlap-markers" in prompt
    assert "repair_overlap" in prompt


def test_segmentation_run_api_lists_runs_and_downloads_exports(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    create_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "export_me.txt",
            "descript_text": "[00:00:00] P: Good morning.\n[00:00:03] Av: Uh yes.",
            "rule_ids": [
                "speaker-markers",
                "timestamp-markers",
                "pause-markers",
                "filled-pauses",
            ],
        },
    )
    run = create_response.json()["run"]

    list_response = client.get("/api/segmentation/runs")
    transcript_response = client.get(
        f"/api/segmentation/runs/{run['run_id']}/exports/final_transcript.txt"
    )
    evidence_response = client.get(
        f"/api/segmentation/runs/{run['run_id']}/exports/evidence.json"
    )
    specialist_response = client.get(
        f"/api/segmentation/runs/{run['run_id']}/specialists/speaker_turn.html"
    )

    assert list_response.status_code == 200
    assert list_response.json()["runs"][0]["run_id"] == run["run_id"]
    assert transcript_response.status_code == 200
    assert transcript_response.text.startswith(
        "Researcher-provided transcript: export_me"
    )
    assert "P: Good morning." in transcript_response.text
    assert evidence_response.status_code == 200
    assert evidence_response.json()["source"] == "researcher_provided"
    evaluation = evidence_response.json()["evaluation"]
    assert evaluation["passed_rule_count"] == evaluation["configured_rule_count"]
    assert "score" not in evaluation
    assert specialist_response.status_code == 200
    assert "Do not rewrite the full transcript" in specialist_response.text


def test_segmentation_corpus_run_api_creates_and_lists_regression_batch(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)

    create_response = client.post(
        "/api/segmentation/corpus-runs",
        json={"seed": 19},
    )
    list_response = client.get("/api/segmentation/corpus-runs")

    assert create_response.status_code == 200
    corpus_run = create_response.json()["corpus_run"]
    assert corpus_run["status"] == "passed"
    assert corpus_run["seed"] == 19
    assert corpus_run["total_case_count"] == 4
    assert corpus_run["regression_fail_count"] == 0
    synthetic_run = client.get(
        f"/api/segmentation/runs/{corpus_run['results'][0]['run_id']}"
    ).json()["run"]
    assert synthetic_run["source"] == "synthetic"
    assert synthetic_run["merged_draft"].startswith("Synthetic run:")
    assert any(
        result["expected_status"] == "failed"
        and result["failed_rule_ids"] == ["official-source-guard"]
        for result in corpus_run["results"]
    )
    assert list_response.status_code == 200
    assert list_response.json()["corpus_runs"][0]["corpus_run_id"] == corpus_run[
        "corpus_run_id"
    ]


def test_segmentation_run_analysis_api_uses_verified_merged_transcript(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    create_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "analysis_me.txt",
            "descript_text": "[00:00:00] P: Good morning.\n[00:00:03] Av: Uh yes.",
            "rule_ids": [
                "speaker-markers",
                "timestamp-markers",
                "pause-markers",
                "filled-pauses",
            ],
        },
    )
    run = create_response.json()["run"]

    analysis_response = client.post(
        f"/api/segmentation/runs/{run['run_id']}/analysis",
        json={},
    )

    assert analysis_response.status_code == 200
    payload = analysis_response.json()
    assert payload["source_filename"] == "analysis_me_segmented.txt"
    assert payload["turn_count"] == 2
    assert [result["metric_id"] for result in payload["results"]] == [
        "base_metrics",
        "lexical_metrics",
        "disfluency_metrics",
    ]
    assert payload["results"][0]["rows"][1]["speaker"] == "participant"
    assert (tmp_path / "runs" / payload["run_id"] / "results.json").exists()


def test_segmentation_run_analysis_api_honors_metric_config(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    create_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "configured_analysis.txt",
            "descript_text": "[00:00:00] P: Good morning.\n[00:00:03] Av: Uh yes.",
            "rule_ids": [
                "speaker-markers",
                "timestamp-markers",
                "pause-markers",
                "filled-pauses",
            ],
        },
    )
    run = create_response.json()["run"]

    analysis_response = client.post(
        f"/api/segmentation/runs/{run['run_id']}/analysis",
        json={
            "config": {
                "selected_metrics": ["base_metrics"],
                "disfluency_tokens": ["yes"],
            }
        },
    )

    assert analysis_response.status_code == 200
    payload = analysis_response.json()
    assert [result["metric_id"] for result in payload["results"]] == ["base_metrics"]
    assert payload["results"][0]["rows"][2]["turns"] == 2


def test_segmentation_run_analysis_api_rejects_unverified_runs(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    create_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "leak.txt",
            "descript_text": "[00:00:00] P: Nala should not appear here.",
            "rule_ids": [
                "speaker-markers",
                "timestamp-markers",
                "official-source-guard",
            ],
        },
    )
    run = create_response.json()["run"]

    analysis_response = client.post(
        f"/api/segmentation/runs/{run['run_id']}/analysis",
        json={},
    )

    assert analysis_response.status_code == 400
    assert analysis_response.json()["detail"] == (
        "Segmentation run must be verified before analysis"
    )


def test_segmentation_run_api_accepts_specialist_patch_submission(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    create_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "patch_me.txt",
            "descript_text": "[00:00:00] P: Good morning.\n[00:00:03] Av: Uh yes.",
            "rule_ids": [
                "speaker-markers",
                "timestamp-markers",
                "pause-markers",
                "filled-pauses",
            ],
        },
    )
    run = create_response.json()["run"]

    patch_response = client.post(
        f"/api/segmentation/runs/{run['run_id']}/specialists/timing_pause/patches",
        json={
            "patches": [
                {
                    "operation": "insert_before_event",
                    "event_index": 0,
                    "text": "-0:00",
                    "reason": "submitted by timing/pause agent",
                }
            ]
        },
    )

    assert patch_response.status_code == 200
    updated = patch_response.json()["run"]
    assert updated["status"] == "needs_rewrite"
    assert updated["failure_routes"][0]["specialist_id"] == "timing_pause"
    assert "; :03" not in updated["merged_draft"]


def test_segmentation_run_api_rejects_invalid_specialist_patch_submission(
    tmp_path,
    monkeypatch,
) -> None:
    monkeypatch.setenv("NLP_SKILL_AGENTS_DATA_DIR", str(tmp_path))
    client = TestClient(app)
    create_response = client.post(
        "/api/segmentation/runs",
        json={
            "source_filename": "patch_me.txt",
            "descript_text": "[00:00:00] P: Good morning.",
            "rule_ids": ["speaker-markers", "timestamp-markers"],
        },
    )
    run = create_response.json()["run"]

    patch_response = client.post(
        f"/api/segmentation/runs/{run['run_id']}/specialists/timing_pause/patches",
        json={
            "patches": [
                {
                    "operation": "insert_before_event",
                    "event_index": 99,
                    "text": "-0:00",
                    "reason": "bad event index",
                }
            ]
        },
    )

    assert patch_response.status_code == 400
    assert patch_response.json()["detail"] == "Patch event_index out of range"
